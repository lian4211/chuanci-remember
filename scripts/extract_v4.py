#!/usr/bin/env python3
"""v4 - 修复可能的循环问题"""
import docx, json, re, os

DOCX = "C:/Users/a1370/Desktop/扫描书籍_章节索引(1).docx"
OUT = "C:/Users/a1370/chuanci-remember/data/wordbook.json"

doc = docx.Document(DOCX)

body = doc.element.body
elements = []
t_idx = p_idx = 0
for child in body:
    tag = child.tag.split('}')[-1]
    if tag == 'tbl':
        elements.append(('table', t_idx))
        t_idx += 1
    elif tag == 'p':
        elements.append(('para', p_idx))
        p_idx += 1

print(f"{len(elements)} 元素 ({t_idx} 表格, {p_idx} 段落)")

# 分Unit
units = {}
cu = None
for typ, idx in elements:
    if typ == 'para':
        t = doc.paragraphs[idx].text.strip()
        m = re.match(r'^Unit\s+(\d+)$', t)
        if m:
            cu = int(m.group(1))
            units[cu] = []
            continue
    if cu is not None:
        units[cu].append((typ, idx))

print(f"{len(units)} 单元")

# 解析
all_units = []
total_words = 0

for uid in sorted(units.keys()):
    elems = units[uid]
    words = []
    seen = set()
    i = 0
    max_iter = len(elems) * 2
    iter_count = 0
    
    while i < len(elems) and iter_count < max_iter:
        iter_count += 1
        typ, idx = elems[i]
        
        if typ != 'table':
            i += 1
            continue
        
        table = doc.tables[idx]
        if not table.rows:
            i += 1
            continue
        
        c00 = table.rows[0].cells[0].text.strip().lower()
        if len(table.rows) > 2 or not re.match(r'^[a-z][a-z\'\-]*$', c00) or len(c00) <= 1 or c00 in seen:
            i += 1
            continue
        
        # 音标
        ph = {}
        if len(table.rows[0].cells) > 1:
            t = table.rows[0].cells[1].text.strip()
            us = re.search(r'美\s*(/[^/]+/)', t)
            uk = re.search(r'英\s*(/[^/]+/)', t)
            if us: ph["us"] = us.group(1)
            if uk: ph["uk"] = uk.group(1)
            if not us and not uk:
                m = re.search(r'(/[^/]+/)', t)
                if m: ph["us"] = m.group(1)
        if not ph: ph = {}
        
        # 找后续段落
        j = i + 1
        paras = []
        while j < len(elems):
            nt, ni = elems[j]
            if nt == 'table':
                ntbl = doc.tables[ni]
                if len(ntbl.rows) <= 2:
                    nw = ntbl.rows[0].cells[0].text.strip().lower()
                    if re.match(r'^[a-z][a-z\'\-]*$', nw) and len(nw) > 1:
                        break
                j += 1
                continue
            else:
                t_ = doc.paragraphs[ni].text.strip()
                if t_:
                    paras.append(t_)
                j += 1
        
        # 解析
        entry = {"english": c00, "phonetic": ph, "definitions": [],
                 "exam_sentences": [], "extensions": [], "memory_aid": "", "phrases": []}
        
        for p in paras:
            if p in ('真题例句', '助记', '拓展', '拓展：'):
                continue
            if p.startswith('参考译文'):
                if entry["exam_sentences"]:
                    entry["exam_sentences"][-1]["translation"] = p.replace('参考译文', '').strip()
                continue
            
            if re.match(r'^(?:n\.|v\.|adj\.|adv\.|prep\.|pron\.|conj\.|art\.|vt\.|vi\.|abbr\.|int\.|num\.|det\.)', p):
                entry["definitions"].append(p)
                continue
            
            if re.match(r'^\d+[.)]\s', p):
                entry["exam_sentences"].append({
                    "sentence": re.sub(r'^\d+[.)]\s*', '', p).strip(),
                    "translation": "", "source": ""
                })
                continue
            
            ym = re.search(r'(\(\d{4}\s*年[^)]*\))', p)
            if ym and entry["exam_sentences"] and not entry["exam_sentences"][-1]["source"]:
                entry["exam_sentences"][-1]["source"] = ym.group(1)
                continue
            
            em = re.match(r'^([a-z][a-z\'\-]+)\s+(n\.|v\.|adj\.|adv\.|prep\.|vt\.|vi\.)', p)
            if em:
                ew = em.group(1).lower()
                if ew != c00:
                    rest = re.sub(r'^\S+\s+', '', p, count=1).strip()
                    entry["extensions"].append({"word": ew, "meaning": rest})
                continue
            
            if '=' in p and len(p) > 8:
                entry["memory_aid"] = (entry["memory_aid"] + ' ' + p).strip()
                continue
        
        seen.add(c00)
        words.append(entry)
        i = j
    
    all_units.append({"id": uid, "title": f"Unit {uid}",
                      "words": [{"id": k+1, **w} for k, w in enumerate(words)]})
    total_words += len(words)
    print(f"Unit {uid}: {len(words)} 词")

for uid in sorted(units.keys()):
    pass

wb = {"name": "考研英语串词记忆", "version": "1.0", "total_units": len(all_units), "units": all_units}
with open(OUT, "w", encoding="utf-8") as f:
    json.dump(wb, f, ensure_ascii=False, indent=2)
print(f"\n总共 {total_words} 词, {len(all_units)} 单元")
