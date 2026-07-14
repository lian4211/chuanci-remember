#!/usr/bin/env python3
"""
Chuanci Remember — docx 解析器 v3（表格驱动）
每个单词条 = 1个表格(单词+音标) + 后续段落(释义/例句/拓展)
"""
import docx, json, re, os

DOCX = "C:/Users/a1370/Desktop/扫描书籍_章节索引(1).docx"
OUT = "C:/Users/a1370/chuanci-remember/data/wordbook.json"

doc = docx.Document(DOCX)

# ============= 步骤1：按文档顺序提取所有元素 =============
# 表格和段落交替排列，我们要按文档顺序处理
# docx 中获取 body 的 xml 顺序
body = doc.element.body
elements = []  # [('table', idx) or ('para', idx)]

t_idx = 0
p_idx = 0
for child in body:
    tag = child.tag.split('}')[-1]
    if tag == 'tbl':
        elements.append(('table', t_idx))
        t_idx += 1
    elif tag == 'p':
        elements.append(('para', p_idx))
        p_idx += 1

print(f"文档元素: {len(elements)} ({sum(1 for e in elements if e[0]=='table')} 表格, {sum(1 for e in elements if e[0]=='para')} 段落)")

# ============= 步骤2：按 Unit 分割 =============
units = {}  # unit_id -> [('table',idx) or ('para',idx)]
current_unit = None

for elem in elements:
    typ, idx = elem
    if typ == 'para':
        p = doc.paragraphs[idx]
        t = p.text.strip()
        m = re.match(r'^Unit\s+(\d+)$', t)
        if m:
            current_unit = int(m.group(1))
            units[current_unit] = []
            continue
    if current_unit is not None:
        units[current_unit].append(elem)

print(f"找到 {len(units)} 个单元")

# ============= 步骤3：解析每个单元的单词 =============
def parse_unit(unit_id):
    """解析一个单元的所有单词"""
    elems = units.get(unit_id, [])
    words = []
    seen_words = set()
    i = 0
    
    while i < len(elems):
        typ, idx = elems[i]
        
        if typ != 'table':
            i += 1
            continue
        
        table = doc.tables[idx]
        if len(table.rows) == 0:
            i += 1
            continue
        
        # 检查表格是否为单词条目（[0,0] 是英文单词）
        cell00 = table.rows[0].cells[0].text.strip()
        
        # 跳过章节索引表（多行，内容为单词链）
        if len(table.rows) > 2:
            i += 1
            continue
        
        # 检查是否为词条（单词+音标）
        word = cell00.lower().strip()
        if not re.match(r'^[a-z][a-z\'\-]*$', word) or len(word) <= 1:
            i += 1
            continue
        
        if word in seen_words:
            i += 1
            continue
        
        # 提取音标
        phonetic = {"us": "", "uk": ""}
        if len(table.rows[0].cells) > 1:
            ph_text = table.rows[0].cells[1].text.strip()
            # 解析音标
            us_m = re.search(r'美\s*(/[^/]+/)', ph_text)
            uk_m = re.search(r'英\s*(/[^/]+/)', ph_text)
            if us_m: phonetic["us"] = us_m.group(1)
            if uk_m: phonetic["uk"] = uk_m.group(1)
            if not us_m and not uk_m:
                # 可能只有一个音标
                ph = re.search(r'(/[^/]+/)', ph_text)
                if ph: phonetic["us"] = ph.group(1)
        if not phonetic["us"] and not phonetic["uk"]:
            phonetic = {}
        
        # 收集后续段落（直到下一个词条表格或空段落）
        j = i + 1
        paras = []
        while j < len(elems):
            next_typ, next_idx = elems[j]
            if next_typ == 'table':
                # 检查下一表格是否是另一个词条
                next_t = doc.tables[next_idx]
                if len(next_t.rows) <= 2:
                    next_word = next_t.rows[0].cells[0].text.strip().lower()
                    if re.match(r'^[a-z][a-z\'\-]*$', next_word) and len(next_word) > 1:
                        break
                # 章节索引表（多行）
                if len(next_t.rows) > 2:
                    j += 1
                    continue
                break
            else:
                p = doc.paragraphs[next_idx].text.strip()
                if p:
                    paras.append(p)
                j += 1
        
        # 解析段落到各字段
        entry = {
            "english": word,
            "phonetic": phonetic,
            "definitions": [],
            "exam_sentences": [],
            "extensions": [],
            "memory_aid": "",
            "phrases": []
        }
        
        for p in paras:
            if p == '真题例句':
                continue
            if p.startswith('参考译文'):
                if entry["exam_sentences"]:
                    s = entry["exam_sentences"][-1]
                    s["translation"] = (s.get("translation", "") + p.replace('参考译文', '').strip()).strip()
                continue
            if p.startswith('拓展') or p.startswith('拓展：'):
                continue
            if p == '助记':
                continue
            if p.startswith('串词'):
                continue
            
            # 定义行 (POS开头)
            if re.match(r'^(?:n\.|v\.|adj\.|adv\.|prep\.|pron\.|conj\.|art\.|vt\.|vi\.|abbr\.|int\.|num\.|det\.|pl\.|sing\.)', p):
                entry["definitions"].append(p)
                continue
            
            # 例句 (数字开头)
            if re.match(r'^\d+[.)]\s', p):
                entry["exam_sentences"].append({
                    "sentence": re.sub(r'^\d+[.)]\s*', '', p).strip(),
                    "translation": "",
                    "source": ""
                })
                continue
            
            # 年份来源
            year_m = re.search(r'(\(\d{4}\s*年[^)]*\))', p)
            if year_m:
                if entry["exam_sentences"] and not entry["exam_sentences"][-1]["source"]:
                    entry["exam_sentences"][-1]["source"] = year_m.group(1)
                continue
            
            # 拓展词 (word POS ...)
            ext_m = re.match(r'^([a-zA-Z][a-zA-Z\'\-]+)\s+(n\.|v\.|adj\.|adv\.|prep\.|vt\.|vi\.)', p)
            if ext_m:
                ew = ext_m.group(1).lower()
                if ew != word:
                    entry["extensions"].append({"word": ew, "meaning": re.sub(r'^\S+\s+', '', p, count=1).strip()})
                continue
            
            # 助记 (含=号)
            if '=' in p and len(p) > 8:
                entry["memory_aid"] = (entry["memory_aid"] + ' ' + p).strip()
                continue
            
            # 短语 (英文 + 中文)
            phrase_m = re.match(r'^([a-zA-Z\s]+?)\s+([\u4e00-\u9fff].+)', p)
            if phrase_m:
                entry["phrases"].append({
                    "phrase": phrase_m.group(1).strip(),
                    "meaning": phrase_m.group(2).strip()
                })
                continue
            
            # 其他（可能是延续的定义或例句翻译）
            if entry["exam_sentences"] and not entry["exam_sentences"][-1].get("translation"):
                entry["exam_sentences"][-1]["translation"] = p
            elif entry["definitions"]:
                # 可能是定义延续
                pass
        
        seen_words.add(word)
        words.append(entry)
        i = j
    
    return words

# ============= 执行 =============
all_units = []
total_words = 0
for uid in sorted(units.keys()):
    words = parse_unit(uid)
    all_units.append({
        "id": uid,
        "title": f"Unit {uid}",
        "words": [{"id": j+1, **w} for j, w in enumerate(words)]
    })
    total_words += len(words)
    print(f"Unit {uid}: {len(words)} 词")

wb = {
    "name": "考研英语串词记忆",
    "version": "1.0",
    "total_units": len(all_units),
    "units": all_units
}

with open(OUT, "w", encoding="utf-8") as f:
    json.dump(wb, f, ensure_ascii=False, indent=2)

print(f"\n总共 {total_words} 词, {len(all_units)} 单元")
