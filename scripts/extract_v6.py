#!/usr/bin/env python3
"""v6 — 正确处理音标与单词在同一单元格的情况"""
import docx, json, re, os

DOCX = "C:/Users/a1370/Desktop/扫描书籍_章节索引(1).docx"
OUT = "C:/Users/a1370/chuanci-remember/data/wordbook.json"

doc = docx.Document(DOCX)
paras = [(p.text.strip(), i) for i, p in enumerate(doc.paragraphs)]
tables = doc.tables

body = doc.element.body
elements = []
ti = pi = 0
for child in body:
    tag = child.tag.split('}')[-1]
    if tag == 'tbl':
        elements.append(('t', ti)); ti += 1
    elif tag == 'p':
        elements.append(('p', pi)); pi += 1

# 分 Unit - 处理重复的 Unit 标记
units = {}  # uid -> [(type, idx)]
cu = None
prev_unit = None

for typ, idx in elements:
    if typ == 'p':
        t = doc.paragraphs[idx].text.strip()
        m = re.match(r'^Unit\s+(\d+)$', t)
        if m:
            uid = int(m.group(1))
            # 第一次遇到这个单元号，初始化
            if uid not in units:
                units[uid] = []
            cu = uid
            continue
    if cu is not None:
        units[cu].append((typ, idx))

print(f"{len(units)} 个单元")

def extract_word_from_cell(cell_text):
    """从单元格文本中提取英文单词"""
    t = cell_text.strip().lower()
    # 直接是单词
    if re.match(r'^[a-z][a-z\'\-]*$', t) and len(t) > 1:
        return t
    # 按行切分
    for line in t.split('\n'):
        l = line.strip()
        # 去掉前导音标: "美/.../" 或 "英/.../"
        l2 = re.sub(r'^[美英]\s*/[^/]*/\s*', '', l).strip()
        if re.match(r'^[a-z][a-z\'\-]*$', l2) and len(l2) > 1:
            return l2
        # 去掉tab
        l3 = l.split('\t')[0].strip()
        if re.match(r'^[a-z][a-z\'\-]*$', l3) and len(l3) > 1:
            return l3
        # 行中第一个英文词
        wm = re.match(r'^([a-z][a-z\'\-]+)', l)
        if wm:
            return wm.group(1)
    return ""

def extract_phonetic_from_cell(cell_text):
    """从单元格中提取音标"""
    t = cell_text.strip()
    ph = {}
    us = re.search(r'美\s*(/[^/]+/)', t)
    uk = re.search(r'英\s*(/[^/]+/)', t)
    if us: ph["us"] = us.group(1)
    if uk: ph["uk"] = uk.group(1)
    if not us and not uk:
        m = re.search(r'(/[^/]+/)', t)
        if m: ph["us"] = m.group(1)
    return ph

all_units = []
total = 0

for uid in sorted(units.keys()):
    elems = units[uid]
    words = []
    seen = set()
    current = None
    
    for typ, idx in elems:
        if typ == 't':
            table = tables[idx]
            if not table.rows: continue
            
            cell0 = table.rows[0].cells[0].text
            word = extract_word_from_cell(cell0)
            if not word or word in seen: continue
            
            # Save previous
            if current: words.append(current)
            
            ph = {}
            if len(table.rows[0].cells) > 1:
                ph = extract_phonetic_from_cell(table.rows[0].cells[1].text)
            
            seen.add(word)
            current = {"english": word, "phonetic": ph, "definitions": [],
                       "exam_sentences": [], "extensions": [], "memory_aid": "", "phrases": []}
        else:
            if not current: continue
            p = doc.paragraphs[idx].text.strip()
            if not p: continue
            e = current
            
            if p in ('真题例句', '助记', '拓展', '拓展：'): continue
            if p.startswith('参考译文'):
                if e["exam_sentences"]:
                    e["exam_sentences"][-1]["translation"] = p.replace('参考译文', '').strip()
                continue
            if re.match(r'^(?:n\.|v\.|adj\.|adv\.|prep\.|pron\.|conj\.|art\.|vt\.|vi\.|abbr\.|int\.|num\.|det\.)', p):
                e["definitions"].append(p); continue
            if re.match(r'^\d+[.)]\s', p):
                e["exam_sentences"].append({"sentence": re.sub(r'^\d+[.)]\s*', '', p).strip(), "translation": "", "source": ""})
                continue
            ym = re.search(r'(\(\d{4}\s*年[^)]*\))', p)
            if ym and e["exam_sentences"] and not e["exam_sentences"][-1]["source"]:
                e["exam_sentences"][-1]["source"] = ym.group(1); continue
            em = re.match(r'^([a-z][a-z\'\-]+)\s+(n\.|v\.|adj\.|adv\.|prep\.|vt\.|vi\.)', p)
            if em:
                ew = em.group(1).lower()
                if ew != word:
                    e["extensions"].append({"word": ew, "meaning": re.sub(r'^\S+\s+', '', p, count=1).strip()})
                continue
            if '=' in p and len(p) > 8:
                e["memory_aid"] = (e["memory_aid"] + ' ' + p).strip(); continue
    
    if current: words.append(current)
    all_units.append({"id": uid, "title": f"Unit {uid}", "words": [{"id": i+1, **w} for i, w in enumerate(words)]})
    total += len(words)
    print(f"Unit {uid}: {len(words)} 词")

with open(OUT, "w", encoding="utf-8") as f:
    json.dump({"name": "考研英语串词记忆", "version": "1.0", "total_units": len(all_units), "units": all_units}, f, ensure_ascii=False, indent=2)
print(f"\n总共 {total} 词")
