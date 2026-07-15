import docx, json, re

idx = docx.Document("C:/Users/a1370/Desktop/单元索引.docx")

# Build all words
all_words = set()
unit_words = {}
cu = None

for p in idx.paragraphs:
    t = p.text.strip()
    m = re.match(r"^Unit\s*(\d+)$", t)
    if m:
        cu = int(m.group(1))
        if cu not in unit_words:
            unit_words[cu] = set()
        continue
    if cu is not None and t and t != "章节索引" and not t.startswith("章节索引"):
        for part in re.split(r"[\t\s]+", t):
            m2 = re.match(r"^[□﹁\- ]*([a-zA-Z][a-zA-Z\'\-]*)", part.strip())
            if m2:
                w = m2.group(1).lower()
                unit_words[cu].add(w)

# Tables
cu = None
body = idx.element.body
elements = []
ti = pi = 0
for child in body:
    tag = child.tag.split("}")[-1]
    if tag == "tbl": elements.append(("t", ti)); ti += 1
    elif tag == "p": elements.append(("p", pi)); pi += 1

cu = None
for typ, idx2 in elements:
    if typ == "p":
        t = idx.paragraphs[idx2].text.strip()
        m = re.match(r"^Unit\s*(\d+)$", t)
        if m: cu = int(m.group(1))
        if cu not in unit_words: unit_words[cu] = set()
        continue
    if cu is not None and typ == "t":
        table = idx.tables[idx2]
        for row in table.rows:
            for cell in row.cells:
                for part in re.split(r"[\t\s]+", cell.text.strip()):
                    m2 = re.match(r"^[□﹁\- ]*([a-zA-Z][a-zA-Z\'\-]*)", part.strip())
                    if m2:
                        w = m2.group(1).lower()
                        if len(w) > 1:
                            unit_words[cu].add(w)

with open("C:/Users/a1370/chuanci-remember/data/index_words.json", "w", encoding="utf-8") as f:
    data = {str(k): sorted(v) for k, v in sorted(unit_words.items())}
    json.dump(data, f, ensure_ascii=False, indent=2)

total = sum(len(v) for v in unit_words.values())
print(f"{len(unit_words)} 单元, {total} 词")
