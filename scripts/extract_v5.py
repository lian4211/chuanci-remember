#!/usr/bin/env python3
"""v5 — 单遍扫描，O(n)"""
import docx, json, re, os

DOCX = "C:/Users/a1370/Desktop/扫描书籍_章节索引(1).docx"
OUT = "C:/Users/a1370/chuanci-remember/data/wordbook.json"

doc = docx.Document(DOCX)

# 构建元素序列: (type, idx_or_text)
# 先收集所有段落和表格的引用
paras = [(p.text.strip(), i) for i, p in enumerate(doc.paragraphs)]
tables = doc.tables

body = doc.element.body
elements = []
ti = pi = 0
for child in body:
    tag = child.tag.split('}')[-1]
    if tag == 'tbl':
        elements.append(('t', ti))
        ti += 1
    elif tag == 'p':
        elements.append(('p', pi))
        pi += 1

print(f"总元素: {len(elements)}")

# 分段，找 Unit
units = {}  # uid -> [(type, idx)]
cu = None
for elem in elements:
    typ, idx = elem
    if typ == 'p':
        t = doc.paragraphs[idx].text.strip()
        m = re.match(r'^Unit\s+(\d+)$', t)
        if m:
            cu = int(m.group(1))
            units[cu] = []
            continue
    if cu is not None:
        units[cu].append(elem)

print(f"找到 {len(units)} 个单元: {sorted(units.keys())}")

# 解析
def parse_unit(uid):
    elems = units[uid]
    words = []
    seen = set()
    
    current_word = None
    current_entry = None
    
    for typ, idx in elems:
        if typ == 't':
            table = tables[idx]
            if not table.rows: continue
            c00 = table.rows[0].cells[0].text.strip().lower()
            c00 = table.rows[0].cells[0].text.strip().lower()
            
            # 从单元格文本中提取单词
            word = ""
            # 尝试：整行就是单词
            if re.match(r'^[a-z][a-z\'\-]*$', c00) and len(c00) > 1:
                word = c00
            else:
                # 尝试：多行中的单词（音标在前，单词在后）
                for line in c00.split('\n'):
                    l = line.strip()
                    # 去掉可能的前导音标
                    l_clean = re.sub(r'^[美英]\s*/[^/]*/\s*', '', l).strip()
                    # 也去掉 tab 后的内容
                    l_clean = l_clean.split('\t')[0].strip()
                    if re.match(r'^[a-z][a-z\'\-]*$', l_clean) and len(l_clean) > 1:
                        word = l_clean
                        break
                    # 直接匹配行中第一个英文单词
                    wm = re.match(r'^([a-z][a-z\'\-]+)', l)
                    if wm:
                        word = wm.group(1)
                        break
            
            if not word or len(word) <= 1 or word in seen:
                continue
                if current_entry:
                    words.append(current_entry)
                
                # 音标
                ph = {}
                if len(table.rows[0].cells) > 1:
                    t = table.rows[0].cells[1].text.strip()
                    us = re.search(r'美\s*(/[^/]+/)', t)
                    uk = re.search(r'英\s*(/[^/]+/)', t)
                    if us: ph["us"] = us.group(1)
                    if uk: ph["uk"] = uk.group(1)
                    if not us and not uk:
                        m = re.search(r'(/[^/]+/)', t)
                        if m: ph["us"] = m.group(1)
                
                current_word = c00
                seen.add(c00)
                current_entry = {
                    "english": c00, "phonetic": ph,
                    "definitions": [], "exam_sentences": [],
                    "extensions": [], "memory_aid": "", "phrases": []
                }
                continue
            else:
                # 章节索引表或多行表，跳过
                continue
        
        else:  # p
            p_text = doc.paragraphs[idx].text.strip()
            if not p_text or not current_entry:
                continue
            
            e = current_entry
            
            if p_text in ('真题例句', '助记', '拓展', '拓展：'):
                continue
            if p_text.startswith('参考译文'):
                if e["exam_sentences"]:
                    e["exam_sentences"][-1]["translation"] = p_text.replace('参考译文', '').strip()
                continue
            
            if re.match(r'^(?:n\.|v\.|adj\.|adv\.|prep\.|pron\.|conj\.|art\.|vt\.|vi\.|abbr\.|int\.|num\.|det\.)', p_text):
                e["definitions"].append(p_text)
                continue
            
            if re.match(r'^\d+[.)]\s', p_text):
                e["exam_sentences"].append({
                    "sentence": re.sub(r'^\d+[.)]\s*', '', p_text).strip(),
                    "translation": "", "source": ""
                })
                continue
            
            ym = re.search(r'(\(\d{4}\s*年[^)]*\))', p_text)
            if ym and e["exam_sentences"] and not e["exam_sentences"][-1]["source"]:
                e["exam_sentences"][-1]["source"] = ym.group(1)
                continue
            
            em = re.match(r'^([a-z][a-z\'\-]+)\s+(n\.|v\.|adj\.|adv\.|prep\.|vt\.|vi\.)', p_text)
            if em:
                ew = em.group(1).lower()
                if ew != current_word:
                    rest = re.sub(r'^\S+\s+', '', p_text, count=1).strip()
                    e["extensions"].append({"word": ew, "meaning": rest})
                continue
            
            if '=' in p_text and len(p_text) > 8:
                e["memory_aid"] = (e["memory_aid"] + ' ' + p_text).strip()
                continue
    
    if current_entry:
        words.append(current_entry)
    
    return words

all_units = []
total = 0
for uid in sorted(units.keys()):
    words = parse_unit(uid)
    wo = [{"id": i+1, **w} for i, w in enumerate(words)]
    all_units.append({"id": uid, "title": f"Unit {uid}", "words": wo})
    total += len(words)
    print(f"Unit {uid}: {len(words)} 词")

with open(OUT, "w", encoding="utf-8") as f:
    json.dump({"name": "考研英语串词记忆", "version": "1.0", "total_units": len(all_units), "units": all_units}, f, ensure_ascii=False, indent=2)
print(f"\n总共 {total} 词")
